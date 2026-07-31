import copy
import re
import shutil
import zipfile
import os
from docx import Document
from docx.oxml.ns import qn
from PIL import Image

FIG_DIR = "/Users/rahul/Desktop/Residency/Research 2/Acute THA vs ORIF for Geriatric Posterior Wall Acetabular Fxs Markov - Meir Marmor/JAAOS_Figures_Tables"
SRC_MANUSCRIPT = "/Users/rahul/Downloads/Markov pTHR 7-26.docx"
SRC_TABLES = FIG_DIR + "/JAAOS_Tables.docx"
WORKDIR = "/tmp/assemble_final"
INTERMEDIATE = WORKDIR + "/intermediate.docx"
FINAL = WORKDIR + "/final.docx"

# =====================================================================
# PART 0: TEXT / XML CORRECTNESS FIXES
# =====================================================================
doc = Document(SRC_MANUSCRIPT)
tdoc = Document(SRC_TABLES)
body = doc.element.body


def para_text(p):
    return ''.join(t.text or '' for t in p.iter(qn('w:t')))


def find_body_para(needle):
    for p in body.iter(qn('w:p')):
        if needle in para_text(p):
            return p
    raise ValueError(f"paragraph containing {needle!r} not found")


def flatten_nested_runs(p):
    """Fix malformed <w:r><w:r>...</w:r></w:r> nesting (invalid OOXML that
    causes some renderers to drop the outer run's font and fall back to the
    style default). Promotes inner runs to direct children of the paragraph,
    inheriting the outer run's rPr when the inner run has none of its own."""
    fixed = 0
    for r in list(p.findall(qn('w:r'))):
        inner_runs = r.findall(qn('w:r'))
        if not inner_runs:
            continue
        outer_rpr = r.find(qn('w:rPr'))
        parent = r.getparent()
        idx = list(parent).index(r)
        for offset, ir in enumerate(inner_runs):
            if ir.find(qn('w:rPr')) is None and outer_rpr is not None:
                ir.insert(0, copy.deepcopy(outer_rpr))
            parent.insert(idx + offset, ir)
        parent.remove(r)
        fixed += len(inner_runs)
    return fixed


def set_paragraph_text(p, new_text):
    """Replace a paragraph's full text with new_text as a single run,
    preserving the paragraph's formatting (taken from its first run's rPr)."""
    runs = p.findall(qn('w:r'))
    assert runs, "paragraph has no runs to preserve formatting from"
    first_rpr = runs[0].find(qn('w:rPr'))
    for r in runs:
        p.remove(r)
    new_r = p.makeelement(qn('w:r'), {})
    if first_rpr is not None:
        new_r.append(copy.deepcopy(first_rpr))
    new_t = p.makeelement(qn('w:t'), {})
    new_t.set(qn('xml:space'), 'preserve')
    new_t.text = new_text
    new_r.append(new_t)
    p.append(new_r)


# --- Fix 1: malformed nested-run XML causing Arial font fallback ---
# NOTE: these two fixes are specific to the 7-26 source manuscript's known
# corruption. If SRC_MANUSCRIPT changes, re-run the audit in
# Model_Audit_2026-07-26/ before assuming these still apply.
p_cycle = find_body_para("A 3-month cycle length was selected")
n1 = flatten_nested_runs(p_cycle)

p_cohort = find_body_para("A modeled cohort of 1,000 patients was evaluated")
n2 = flatten_nested_runs(p_cohort)
print(f"[0a] Flattened malformed nested runs: {n1} + {n2} = {n1 + n2} (expect 6)")

# --- Fix 2: stale Table 1/2/3 cross-reference (Table 1 is now occupancy,
# not events; costs+QALYs are combined into Table 3, not split across 2/3) ---
set_paragraph_text(
    p_cohort,
    "A modeled cohort of 1,000 patients was evaluated in each treatment arm "
    "across 90-day, 1-year, 2-year, 5-year, and 10-year horizons. "
    "End-of-horizon state occupancy is summarized in Table 1, cumulative "
    "event profiles in Table 2, and cost and QALY outcomes in Table 3. "
    "For interpretability, cohort totals can also be expressed on a "
    "per-patient basis by dividing each cumulative value by 1,000."
)
print("[0b] Fixed stale Table 1/2/3 cross-reference.")

# --- Fix 3: "Early outcomes" 90-day figures contradicted the verified
# Table 1 data (THA stable-well state was 84.31%, not 90.39%; THA 90-day
# mortality was 9.11%, not 3.04% -- and is in fact *higher* than ORIF's
# 8.42%, not lower as the old sentence claimed). Verified against
# build_tables.py Table 1/Table 2 source rows (pandas-computed from the
# model's own CSV export). ---
p_early = find_body_para("Early outcomes favored acute THA")
set_paragraph_text(
    p_early,
    "Early outcomes favored acute THA on most measures. At 90 days, 84.31% "
    "of the acute THA cohort occupied a stable well state compared with "
    "45.82% of the ORIF cohort, and medical complications (3.48% vs "
    "35.50%) and infection-related events (1.43% vs 4.26%) were both "
    "higher after ORIF; 90-day mortality was slightly higher after acute "
    "THA than ORIF (9.11% vs 8.42%). Ninety-day cumulative costs were "
    "lower for acute THA and cumulative QALYs were higher (Tables 1–3)."
)
print("[0c] Fixed 90-day outcome figures to match verified Table 1/2 data.")

# =====================================================================
# PART 1: TABLE REPLACEMENT (7 main/supplemental + 2 appendix)
# =====================================================================
tdoc_body_children = list(tdoc.element.body)
TDOC_TBL_POSITIONS = [i for i, c in enumerate(tdoc_body_children) if c.tag == qn('w:tbl')]

# ---------------------------------------------------------------------
# Section-break plumbing. copy_table_and_footnote() only ever copies the
# <w:tbl> + footnote <w:p> -- it does NOT carry over section/orientation,
# so every "landscape" supplemental/appendix table was silently landing in
# the manuscript's default portrait section and overflowing both margins
# (this was the root cause of the cut-off/wrapped appendix tables). tdoc
# itself has real section breaks (see below); we replicate that structure
# around the inserted S1-A2 block: close the manuscript's portrait section
# before S1, reopen it after A2's footnote. Do NOT drop this step even if
# it looks redundant -- table content alone carries no orientation.
# ---------------------------------------------------------------------
PORTRAIT_SECTPR_TEMPLATE = copy.deepcopy(list(body)[-1])
assert PORTRAIT_SECTPR_TEMPLATE.tag == qn('w:sectPr'), "expected trailing body sectPr"


def find_landscape_sectPr(tdoc_body):
    for c in tdoc_body:
        if c.tag != qn('w:p'):
            continue
        pPr = c.find(qn('w:pPr'))
        if pPr is None:
            continue
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is None:
            continue
        pgSz = sectPr.find(qn('w:pgSz'))
        if pgSz is not None and pgSz.get(qn('w:orient')) == 'landscape':
            return sectPr
    raise ValueError("no landscape sectPr found in tdoc")


LANDSCAPE_SECTPR_TEMPLATE = copy.deepcopy(find_landscape_sectPr(tdoc_body_children))


def section_break_para(sectPr_template):
    """A blank paragraph carrying a section-break sectPr, detached so it can
    be positioned with insert_before(). NOTE: the caller must not ALSO add
    an explicit page_break_para() immediately adjacent to this -- a section
    break already forces a new page, and stacking both produces a spurious
    blank page (found and fixed during the 2026-07-30 formatting audit)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pPr.append(copy.deepcopy(sectPr_template))
    el = p._p
    el.getparent().remove(el)
    return el


def txt(c):
    return ''.join(t.text or '' for t in c.iter(qn('w:t')))


def children():
    return list(body)


def find_para(needle, occurrence=1):
    hits = [c for c in children() if c.tag == qn('w:p') and needle in txt(c)]
    return hits[occurrence - 1]


def remove_between(start_el, end_el, inclusive=True):
    ch = children()
    si = ch.index(start_el)
    ei = ch.index(end_el)
    assert si <= ei, f"start {si} after end {ei}"
    lo, hi = (si, ei) if inclusive else (si + 1, ei - 1)
    for el in ch[lo:hi + 1]:
        body.remove(el)


def page_break_para():
    p = doc.add_paragraph()
    run = p.add_run()
    br = run._r.makeelement(qn('w:br'), {})
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    el = p._p
    el.getparent().remove(el)
    return el


def copy_table_and_footnote(nth_table):
    pos = TDOC_TBL_POSITIONS[nth_table]
    tbl_el = tdoc_body_children[pos]
    foot_el = tdoc_body_children[pos + 1]
    assert foot_el.tag == qn('w:p'), f"expected footnote paragraph after table {nth_table}"
    return copy.deepcopy(tbl_el), copy.deepcopy(foot_el)


def insert_before(anchor, *elements):
    for el in elements:
        anchor.addprevious(el)


# --- 1. Main-text Tables 1/2/3 (portrait) ---
cap1 = find_para("Table 1. Cumulative modeled events")
cap3 = find_para("Table 3. Cumulative quality-adjusted life-years")
ch = children()
old_tbl3 = ch[ch.index(cap3) + 1]
assert old_tbl3.tag == qn('w:tbl')

t1, f1 = copy_table_and_footnote(0)
t2, f2 = copy_table_and_footnote(1)
t3, f3 = copy_table_and_footnote(2)

insert_before(cap1, page_break_para(), t1, f1)
insert_before(cap1, page_break_para(), t2, f2)
insert_before(cap1, page_break_para(), t3, f3)

remove_between(cap1, old_tbl3, inclusive=True)
print("[1/9] Main-text Tables 1/2/3 replaced.")

# --- 2. Supplemental Table S1 (opens the landscape island) ---
capS1 = find_para("Supplemental Table S1. Transition probabilities")
capS2 = find_para("Supplemental Table S2. Direct cost inputs")
ch = children()
j = ch.index(capS2) - 1
while ch[j].tag != qn('w:tbl'):
    j -= 1
old_tblS1_last = ch[j]

s1, fs1 = copy_table_and_footnote(3)
insert_before(capS1, section_break_para(PORTRAIT_SECTPR_TEMPLATE), s1, fs1)
remove_between(capS1, old_tblS1_last, inclusive=True)
print("[2/9] Supplemental Table S1 replaced (portrait section closed; landscape island opened).")

# --- 3. Supplemental Table S2 (landscape) ---
capS2 = find_para("Supplemental Table S2. Direct cost inputs")
capS3 = find_para("Supplemental Table S3. Health-state utility")
ch = children()
j = ch.index(capS3) - 1
while ch[j].tag != qn('w:tbl'):
    j -= 1
old_tblS2 = ch[j]

s2, fs2 = copy_table_and_footnote(4)
insert_before(capS2, page_break_para(), s2, fs2)
remove_between(capS2, old_tblS2, inclusive=True)
print("[3/9] Supplemental Table S2 replaced.")

# --- 4. Supplemental Table S3 (+ new S4 appended) (landscape) ---
capS3 = find_para("Supplemental Table S3. Health-state utility")
capAppendix = find_para("Supplementary Material")
ch = children()
j = ch.index(capAppendix) - 1
while ch[j].tag != qn('w:tbl'):
    j -= 1
old_tblS3 = ch[j]

s3, fs3 = copy_table_and_footnote(5)
s4, fs4 = copy_table_and_footnote(6)
insert_before(capS3, page_break_para(), s3, fs3)
insert_before(capS3, page_break_para(), s4, fs4)
remove_between(capS3, old_tblS3, inclusive=True)
print("[4/9] Supplemental Table S3 replaced, S4 appended.")

# --- 5. Appendix A1. State list -> Appendix A1. Model Health States (landscape) ---
capA1 = find_para("Appendix A1. State list")
ch = children()
old_tblA1 = ch[ch.index(capA1) + 1]
assert old_tblA1.tag == qn('w:tbl'), "expected table immediately after Appendix A1 caption"

a1, fa1 = copy_table_and_footnote(7)
insert_before(capA1, page_break_para(), a1, fa1)
remove_between(capA1, old_tblA1, inclusive=True)
print("[5/9] Appendix A1 replaced.")

# --- 6. Appendix A2. Allowed transition logic (closes the landscape island) ---
capA2 = find_para("Appendix A2. Allowed transition logic")
ch = children()
old_tblA2 = ch[ch.index(capA2) + 1]
assert old_tblA2.tag == qn('w:tbl'), "expected table immediately after Appendix A2 caption"

a2, fa2 = copy_table_and_footnote(8)
insert_before(capA2, page_break_para(), a2, fa2, section_break_para(LANDSCAPE_SECTPR_TEMPLATE))
remove_between(capA2, old_tblA2, inclusive=True)
print("[6/9] Appendix A2 replaced (landscape island closed; portrait resumes via original trailing sectPr).")

doc.save(INTERMEDIATE)
print("Saved intermediate:", INTERMEDIATE)
print("Final table count:", len(Document(INTERMEDIATE).tables), "(expect 9)")

# =====================================================================
# PART 2: IMAGE SWAP (direct media-byte replacement) + EXTENT FIX
# =====================================================================
UNZIP_DIR = WORKDIR + "/docx_unzipped"
if os.path.exists(UNZIP_DIR):
    shutil.rmtree(UNZIP_DIR)
os.makedirs(UNZIP_DIR)

with zipfile.ZipFile(INTERMEDIATE) as z:
    z.extractall(UNZIP_DIR)

# rId -> (media filename, new PNG source). Confirmed via word/_rels/document.xml.rels;
# re-verify this mapping if SRC_MANUSCRIPT is regenerated from Word (rIds can shift).
IMAGE_MAP = {
    "rId8": ("image1.png", FIG_DIR + "/Figure1_StateMap.png"),
    "rId11": ("image4.png", FIG_DIR + "/Figure2_OneWaySensitivity.png"),
    "rId12": ("image5.png", FIG_DIR + "/Figure3_TwoWayUtilities.png"),
    "rId13": ("image6.png", FIG_DIR + "/Figure4_TwoWayCosts.png"),
}

# Verify the rels still point where we expect before overwriting anything.
rels_path = UNZIP_DIR + "/word/_rels/document.xml.rels"
rels_xml = open(rels_path, encoding="utf-8").read()
for rid, (media_name, _) in IMAGE_MAP.items():
    assert f'Id="{rid}"' in rels_xml and f'media/{media_name}"' in rels_xml, \
        f"rel mapping changed for {rid}; expected media/{media_name}"

# Overwrite media bytes.
for rid, (media_name, new_png) in IMAGE_MAP.items():
    dst = f"{UNZIP_DIR}/word/media/{media_name}"
    shutil.copyfile(new_png, dst)
print("[7/9] Media bytes swapped for all 4 figures.")

# Recompute extents: keep existing cx (width), recompute cy from new aspect ratio,
# so a figure regenerated at a different aspect ratio never stretches on insert.
doc_xml_path = UNZIP_DIR + "/word/document.xml"
xml = open(doc_xml_path, encoding="utf-8").read()

for rid, (media_name, new_png) in IMAGE_MAP.items():
    im = Image.open(new_png)
    new_aspect = im.size[0] / im.size[1]

    idx = xml.find(f'r:embed="{rid}"')
    assert idx != -1, f"{rid} not found in document.xml"
    window = xml[max(0, idx - 1500):idx]
    matches = list(re.finditer(r'<wp:extent cx="(\d+)" cy="(\d+)"', window))
    assert matches, f"no wp:extent found before {rid}"
    cx = int(matches[-1].group(1))
    old_cy = int(matches[-1].group(2))
    new_cy = round(cx / new_aspect)

    old_pair = f'cx="{cx}" cy="{old_cy}"'
    new_pair = f'cx="{cx}" cy="{new_cy}"'
    count = xml.count(old_pair)
    assert count == 2, f"{rid}: expected 2 occurrences (wp:extent + a:ext) of '{old_pair}', found {count}"
    xml = xml.replace(old_pair, new_pair)
    print(f"    {rid} ({media_name}): cx={cx} cy {old_cy} -> {new_cy} (aspect {new_aspect:.4f})")

with open(doc_xml_path, "w", encoding="utf-8") as f:
    f.write(xml)
print("[8/9] Image extents updated to match new aspect ratios (no stretching).")

# Re-zip to FINAL.
if os.path.exists(FINAL):
    os.remove(FINAL)
with zipfile.ZipFile(FINAL, "w", zipfile.ZIP_DEFLATED) as zf:
    for root, _, files in os.walk(UNZIP_DIR):
        for fn in files:
            full = os.path.join(root, fn)
            arc = os.path.relpath(full, UNZIP_DIR)
            zf.write(full, arc)
print("[9/9] Saved final:", FINAL)

# Sanity: re-open with python-docx to confirm it's not corrupt.
check = Document(FINAL)
print("Re-open check OK. Table count:", len(check.tables), "Paragraph count:", len(check.paragraphs))
