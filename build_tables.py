#!/usr/bin/env python3
"""Build all JAAOS tables (main-text 1-3, supplemental S1-S4, appendix A1-A2) in
JBJS/Acuna house style, using the corrected post-mortality-fix model output as the
sole numeric source (Model_Audit_2026-07-26/matt_repo_fixed, Summary_Outputs.csv /
Event_Outputs.csv). Gray title banner, matching the confirmed reference example
(Sam Alfonso / Acuna JBJS table) and the global house-style doc -- not black, not navy."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "JAAOS_Tables.docx")
BANNER = "808080"  # gray, matching the confirmed reference table
WHITE = RGBColor(0xFF, 0xFF, 0xFF); GREY = RGBColor(0x60, 0x60, 0x60)
CTR = WD_ALIGN_PARAGRAPH.CENTER; LFT = WD_ALIGN_PARAGRAPH.LEFT

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Times New Roman"; st.font.size = Pt(11)

def set_landscape(section):
    # New sections inherit the previous section's (possibly already-landscape) page
    # dimensions, so only swap w/h if the page is currently taller than it is wide --
    # swapping unconditionally double-flips an inherited landscape section back to a
    # portrait-shaped page while leaving the orientation flag saying "landscape".
    w, h = section.page_width, section.page_height
    if h > w:
        section.page_width, section.page_height = h, w
    section.orientation = WD_ORIENT.LANDSCAPE
    section.top_margin = Inches(0.5); section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6); section.right_margin = Inches(0.6)

def _shade(cell, hx):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hx); tcPr.append(shd)

def _border(cell, edge, sz, color="000000"):
    tcPr = cell._tc.get_or_add_tcPr(); tb = tcPr.find(qn('w:tcBorders'))
    if tb is None: tb = OxmlElement('w:tcBorders'); tcPr.append(tb)
    e = tb.find(qn('w:' + edge))
    if e is None: e = OxmlElement('w:' + edge); tb.append(e)
    e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz)); e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)

def _fill(cell, text, bold=False, italic=False, size=9, color=None, align=LFT):
    cell.text = ""; p = cell.paragraphs[0]; p.alignment = align
    p.paragraph_format.space_before = Pt(0.5); p.paragraph_format.space_after = Pt(0.5)
    p.paragraph_format.line_spacing = 0.98
    r = p.add_run(str(text)); r.bold = bold; r.italic = italic; r.font.size = Pt(size); r.font.name = "Arial"
    if color is not None: r.font.color.rgb = color
    return p

def _widths(t, widths):
    t.autofit = False; t.allow_autofit = False
    tblPr = t._tbl.tblPr; lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'), 'fixed'); tblPr.append(lay)
    for gc, w in zip(t._tbl.tblGrid.findall(qn('w:gridCol')), widths):
        gc.set(qn('w:w'), str(int(w * 1440)))

def jbjs_table(title, headers, body, footnote, span=None, spans=None, widths=None, hsize=9, bsize=9, page_break=True):
    """span: single (text, c0, c1) tuple. spans: list of (text, c0, c1) for multiple
    spanning group headers on one row (e.g., 'Acute THA' over cols 1-2, 'ORIF' over 3-4)."""
    if page_break: doc.add_page_break()
    n = len(headers)
    t = doc.add_table(rows=1, cols=n); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if widths: _widths(t, widths)
    b = t.rows[0].cells[0]
    for j in range(1, n): b = b.merge(t.rows[0].cells[j])
    _shade(b, BANNER); _fill(b, title, bold=True, size=10, color=WHITE, align=LFT)
    all_spans = spans if spans else ([span] if span else [])
    if all_spans:
        sr = t.add_row()
        for txt, c0, c1 in all_spans:
            m = sr.cells[c0]
            for j in range(c0 + 1, c1 + 1): m = m.merge(sr.cells[j])
            _fill(m, txt, bold=True, size=hsize, align=CTR); _border(m, 'bottom', 4)
    hr = t.add_row()
    for j, h in enumerate(headers):
        _fill(hr.cells[j], h, bold=True, size=hsize, align=LFT if j == 0 else CTR); _border(hr.cells[j], 'bottom', 8)
    for item in body:
        row = t.add_row()
        if item[0] == "sub":
            _fill(row.cells[0], item[1], bold=True, italic=True, size=bsize)
            for j in range(1, n): _fill(row.cells[j], "", size=bsize)
        else:
            vals = item[1]; indent = item[2]
            for j, v in enumerate(vals):
                p = _fill(row.cells[j], v, size=bsize, align=LFT if j == 0 else CTR)
                if j == 0 and indent: p.paragraph_format.left_indent = Inches(0.14)
    for c in t.rows[-1].cells: _border(c, 'bottom', 8)
    para(footnote, italic=True, size=7.5, color=GREY)

def para(t, italic=False, bold=False, size=9, align=None, color=None):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic = italic; r.bold = bold; r.font.size = Pt(size); r.font.name = "Arial"
    if color: r.font.color.rgb = color
    if align is not None: p.alignment = align
    return p

# =========================================================================
# TABLE 1 — End-of-horizon state occupancy (Rohit item 4: separate from events)
# =========================================================================
occ_rows = [
    ("90 days",  "843.1 (84.3)", "91.1 (9.1)",  "458.2 (45.8)", "84.2 (8.4)"),
    ("1 year",   "781.0 (78.1)", "150.4 (15.0)","768.7 (76.9)", "182.6 (18.3)"),
    ("2 years",  "776.4 (77.6)", "166.4 (16.6)","750.7 (75.1)", "197.4 (19.7)"),
    ("5 years",  "685.7 (68.6)", "212.0 (21.2)","712.2 (71.2)", "240.1 (24.0)"),
    ("10 years", "529.5 (53.0)", "318.5 (31.8)","610.2 (61.0)", "340.7 (34.1)"),
]
jbjs_table(
    "Table 1.  End-of-Horizon Health-State Occupancy per 1,000 Simulated Patients",
    ["Time Horizon", "Stable well state, n (%)", "Death, n (%)", "Stable well state, n (%)", "Death, n (%)"],
    [("row", list(r), False) for r in occ_rows],
    "The values are given as the expected number of the 1,000-patient cohort occupying the state at "
    "the end of the horizon, with the percentage of the starting cohort in parentheses. These are "
    "terminal-state proportions, not cumulative event counts; see Table 2 for cumulative events. "
    "THA = total hip arthroplasty, ORIF = open reduction and internal fixation.",
    spans=[("Acute THA (n = 1,000)", 1, 2), ("ORIF (n = 1,000)", 3, 4)],
    widths=[1.35, 1.7, 1.35, 1.7, 1.35], page_break=False)

# =========================================================================
# TABLE 2 — Cumulative modeled events per 1,000 patients (90-day and 10-year headline horizons)
# =========================================================================
ev_rows = [
    ("Medical complication", "34.8 (3.5)", "355.0 (35.5)", "121.0 (12.1)", "355.0 (35.5)"),
    ("Infection-related event", "14.3 (1.4)", "42.6 (4.3)", "47.0 (4.7)", "81.6 (8.2)"),
    ("Dislocation event", "14.6 (1.5)", "6.0 (0.6)", "285.3 (28.5)", "21.4 (2.1)"),
    ("Revision event", "0.0 (0.0)", "0.0 (0.0)", "152.2 (15.2)", "6.0 (0.6)"),
    ("Conversion to THA", "0.0 (0.0)", "16.5 (1.6)", "0.0 (0.0)", "53.1 (5.3)"),
    ("ORIF reoperation", "0.0 (0.0)", "35.0 (3.5)", "0.0 (0.0)", "235.3 (23.5)"),
    ("Refracture / periprosthetic fracture", "0.6 (0.1)", "2.5 (0.2)", "15.2 (1.5)", "8.2 (0.8)"),
]
jbjs_table(
    "Table 2.  Cumulative Modeled Events per 1,000 Simulated Patients",
    ["Event", "THA, n (%)", "ORIF, n (%)", "THA, n (%)", "ORIF, n (%)"],
    [("row", [r[0], r[1], r[2], r[3], r[4]], False) for r in ev_rows],
    "The values are given as the expected cumulative number of events per 1,000 patients from model "
    "start through the stated horizon, with the percentage of the starting cohort in parentheses. "
    "These are recurrent expected event counts, not the number of unique patients experiencing at "
    "least one event, and a patient may be represented in more than one event category or more than "
    "once within a category over time. One-year, 2-year, and 5-year event detail is provided in "
    "Supplemental Table S4. THA = total hip arthroplasty, ORIF = open reduction and internal fixation.",
    spans=[("90 Days", 1, 2), ("10 Years", 3, 4)], widths=[2.3, 1.3, 1.3, 1.3, 1.3])

# =========================================================================
# TABLE 3 — Per-patient economic outcomes (Rohit item 5: combined cost+QALY table)
# =========================================================================
econ_rows = [
    ("90 days",  "$42,294", "$56,370", "−$14,076", "0.203", "0.185", "+0.018", "THA dominant"),
    ("1 year",   "$44,698", "$60,535", "−$15,837", "0.731", "0.638", "+0.093", "THA dominant"),
    ("2 years",  "$45,808", "$63,463", "−$17,656", "1.406", "1.233", "+0.173", "THA dominant"),
    ("5 years",  "$49,529", "$65,798", "−$16,270", "3.357", "2.973", "+0.385", "THA dominant"),
    ("10 years", "$57,092", "$69,445", "−$12,352", "6.280", "5.608", "+0.673", "THA dominant"),
]
jbjs_table(
    "Table 3.  Per-Patient Cost-Utility Outcomes, Acute THA versus ORIF",
    ["Time Horizon", "THA Cost", "ORIF Cost", "Incremental Cost*", "THA QALYs", "ORIF QALYs", "Incremental QALYs*", "Result"],
    [("row", list(r), False) for r in econ_rows],
    "*Incremental values are acute THA minus ORIF; a negative incremental cost indicates acute THA "
    "was less costly. Acute THA was less costly and more effective than ORIF at every modeled "
    "horizon (dominant), so no incremental cost-effectiveness ratio is reported. The 10-year "
    "incremental net monetary benefit at a willingness-to-pay threshold of $50,000/QALY was "
    "$45,982,925 for the 1,000-patient cohort ($45,983 per patient). THA = total hip arthroplasty, "
    "ORIF = open reduction and internal fixation, QALY = quality-adjusted life-year.",
    span=None, widths=[1.1, 0.95, 0.95, 1.1, 0.95, 0.95, 1.15, 1.05])

# =========================================================================
# SUPPLEMENTAL TABLE S1 — Transition probabilities (landscape, condensed)
# Quarter-split dislocation rows (Q1-Q4) collapsed to one row each; the
# embedded "<br>" text found in the source table (an HTML-entity artifact,
# not an intentional line break) is replaced with semicolon separators.
# =========================================================================
sec = doc.sections[-1]
new_sec = doc.add_section()
set_landscape(new_sec)

s1_body = [
    ("sub", "Index Surgery Mortality"),
    ("row", ["ORIF", "DEATH", "30-day mortality", "≈8%", "Ljungdahl et al., 2025¹⁹"], False),
    ("row", ["ORIF", "DEATH", "1-year all-cause mortality", "21% (95% CI, 15-26)", "Ljungdahl et al., 2025"], False),
    ("row", ["THA", "DEATH", "30-day mortality", "2.6% (CI, 0.7-4.6)", "Albrektsson et al., 2023²⁰"], False),
    ("row", ["THA", "DEATH", "3-month mortality", "8.7% (CI, 5.2-12.0)", "Albrektsson et al., 2023"], False),
    ("row", ["THA", "DEATH", "1-year mortality", "14.4% (CI, 10.0-18.7)", "Albrektsson et al., 2023"], False),
    ("sub", "ORIF Complications"),
    ("row", ["ORIF", "Medical complication", "90-day medical complications (operative subgroup ≥60y)", "35.5% (37% overall)", "Simske et al., 2023¹"], False),
    ("row", ["ORIF well", "Early SSI", "SSI (deep + superficial), ≤4-6 wk", "4.26% (quarter 1; 82% of 5.2% total)", "Suzuki et al., 2010²¹"], False),
    ("row", ["ORIF well", "Late SSI", "SSI (late), 6 wk-1 y", "0.31% per quarter (quarters 2-4; 18% of 5.2%)", "Suzuki et al., 2010"], False),
    ("row", ["ORIF well", "Aseptic failure", "Aseptic reoperation, year 1", "≈7% total (quarterly: 3.5%, 2.0%, 1.0%, 0.5%)", "Ljungdahl et al., 2025¹⁹"], False),
    ("row", ["ORIF well", "Aseptic failure", "Aseptic reoperation, year 2", "≈5% total (quarterly: 2.0%, 1.5%, 1.0%, 0.5%)", "Ljungdahl et al., 2025"], False),
    ("row", ["ORIF well", "Refracture", "Re-fracture after acetabular ORIF (≥65y)", "≤1% over 1-2 y (0.25% per quarter, year 1)", "Lundin et al., 2023⁴²"], False),
    ("row", ["ORIF well", "Native instability", "Native-hip instability / dislocation, year 1", "≤1% total (quarterly: 0.6%, 0.25%, 0.10%, 0.05%)", "Smakaj et al., 2025²²"], False),
    ("row", ["ORIF reoperation", "DEATH", "30-day mortality (non-arthroplasty reoperation)", "0.5%", "Assumption; no directly matching published rate identified"], False),
    ("sub", "ORIF → THA Conversion"),
    ("row", ["ORIF well", "Conversion THA", "Conversion to THA, posterior-wall-specific cohort (the model's calibration target)", "5% by 2 y; timing 52% year 1, 27% year 2", "Firoozabadi et al., 2018¹⁵"], False),
    ("sub", "Primary THA Complications"),
    ("row", ["THA", "Medical complication", "1-year medical complication (medical-only, excludes surgical)", "13.2% within 1 year", "Kelly et al., 2023²³"], False),
    ("row", ["THA well", "Dislocation", "THA dislocation, year 1 (quarterly)", "2.8% total (quarterly: 1.46%, 0.53%, 0.40%, 0.41%)", "Hermansen et al., 2022²⁴"], False),
    ("row", ["THA well", "Dislocation", "THA dislocation, years 2-10 (constant hazard)", "≈11% by 10 y (≈0.29% per quarter)", "Morison et al., 2016²⁸"], False),
    ("row", ["THA well", "PJI", "THA early PJI, ≤90 days", "1.43% (quarter 1)", "Zeng et al., 2023²⁵"], False),
    ("row", ["THA well", "Periprosthetic fracture", "Periprosthetic femoral fracture after THA (fracture-indication cohort)", "2.19 per 1,000 prosthesis-years (≈2.2% by 10 y, constant hazard)", "Lamb et al., 2024²⁶"], False),
    ("row", ["THA well", "Aseptic loosening", "Aseptic acetabular loosening requiring revision (post-traumatic arthritis)", "≈3% by 5 y (≈0.15% per quarter)", "Ranawat et al., 2009²⁷"], False),
]
jbjs_table(
    "Supplemental Table S1.  Transition Probabilities Used in the Markov Model",
    ["From State", "To State", "Event / Transition", "Base-Case Value", "Source"],
    s1_body,
    "Values are per 3-month cycle unless otherwise specified. CI = confidence interval, "
    "ORIF = open reduction and internal fixation, PJI = periprosthetic joint infection, "
    "SSI = surgical site infection, THA = total hip arthroplasty. The ORIF→THA conversion schedule "
    "used in the model is calibrated to the posterior-wall-specific 2-year figure (Firoozabadi et "
    "al., 2018); the general-cohort 5-year/9-year figures are shown for external context only and are "
    "not the calibration target.",
    widths=[1.1, 1.1, 3.0, 2.5, 1.5], hsize=8.5, bsize=8, page_break=False)

# =========================================================================
# SUPPLEMENTAL TABLE S2 — Direct cost inputs
# =========================================================================
new_sec2 = doc.add_section(); set_landscape(new_sec2)
s2_rows = [
    ("Index acute THA", "$40,994", "Kelley et al., 2025¹³"),
    ("Index ORIF", "$47,207", "Kelley et al., 2025¹³"),
    ("Aseptic revision", "$30,224", "Hammat et al., 2025⁴⁰"),
    ("Septic revision", "$57,264", "Hammat et al., 2025⁴⁰"),
    ("Dislocation episode", "$25,256", "Hammat et al., 2025⁴⁰"),
    ("Periprosthetic joint infection episode", "$28,904", "Szymski et al., 2024²⁹ (European cost study; no matching U.S.-specific estimate identified)"),
    ("Conversion to THA", "$51,243", "Hunter et al., 2023⁴³"),
    ("ORIF reoperation", "$47,207", "Kelley et al., 2025¹³"),
    ("Medical complication episode", "$14,882", "Schwarzkopf et al., 2019³⁰"),
]
jbjs_table(
    "Supplemental Table S2.  Direct Cost Inputs Used in the Model",
    ["Cost Parameter", "Base-Case Value (USD)", "Source"],
    [("row", list(r), False) for r in s2_rows],
    "Costs are used as reported in their original source publications (spanning 2019-2025) without "
    "adjustment for inflation to a common reference year. All costs are U.S. estimates except the "
    "periprosthetic joint infection cost, sourced from a European systematic review. "
    "THA = total hip arthroplasty, ORIF = open reduction and internal fixation.",
    widths=[3.2, 2.0, 4.0], hsize=9, bsize=9, page_break=False)

# =========================================================================
# SUPPLEMENTAL TABLE S3 — Health-state utilities
# =========================================================================
s3_rows = [
    ("Well, ORIF (native hip)", "0.74", "Rommens et al., 2020³¹; ICUROS (Talevski et al., 2021)³²"),
    ("Well, THA", "0.81", "Gordon et al., 2014³³"),
    ("Dislocation, THA", "0.78", "Hermansen et al., 2022²⁴"),
    ("Dislocation, ORIF (native hip)", "0.78", "Proxy: Hermansen et al., 2022"),
    ("Early PJI", "0.30", "Wildeman et al., 2021³⁴; Walter et al., 2021⁴⁴"),
    ("Late PJI", "0.30", "Wildeman et al., 2021³⁴; Walter et al., 2021⁴⁴"),
    ("Aseptic revision (perioperative)", "0.35", "No direct utility identified; perioperative assumption"),
    ("Septic revision (perioperative)", "0.35", "No direct utility identified; perioperative assumption"),
    ("Post-aseptic revision", "0.70", "Pavlovic et al., 2024³⁵"),
    ("Post-septic revision", "0.70", "Wildeman et al., 2021³⁴; Pavlovic et al., 2024"),
    ("Well, revised", "0.70", "Pavlovic et al., 2024"),
    ("Conversion THA (perioperative)", "0.30", "No direct utility identified; perioperative assumption"),
    ("ORIF reoperation (perioperative)", "0.35", "No direct utility identified; perioperative assumption"),
    ("Post-ORIF reoperation", "0.70", "No direct utility identified; aligned to post-revision literature"),
    ("Aseptic loosening (transient)", "0.35", "No direct utility identified; transient detection state before revision"),
    ("Refracture (transient)", "0.35", "No direct utility identified; transient assumption"),
    ("Medical complication", "0.50", "Loggers et al., 2022³⁶"),
    ("Death", "0.00", "By definition"),
]
jbjs_table(
    "Supplemental Table S3.  Health-State Utility Values Used in the Model",
    ["Health State", "Base-Case Utility", "Source"],
    [("row", list(r), False) for r in s3_rows],
    "Utilities were drawn from published EQ-5D and SF-6D data in hip arthroplasty and fracture "
    "populations; where a direct estimate for a modeled temporary state was unavailable, a proxy "
    "value or short perioperative disutility was assigned and is labeled as such. "
    "PJI = periprosthetic joint infection, THA = total hip arthroplasty, ORIF = open reduction and "
    "internal fixation.",
    widths=[3.5, 1.7, 4.0], hsize=9, bsize=9)

# =========================================================================
# SUPPLEMENTAL TABLE S4 — Full 5-horizon cumulative event detail
# =========================================================================
s4_data = [
    ("90 days", "Medical complication", "34.8 (3.5)", "355.0 (35.5)"),
    ("90 days", "Infection-related event", "14.3 (1.4)", "42.6 (4.3)"),
    ("90 days", "Dislocation event", "14.6 (1.5)", "6.0 (0.6)"),
    ("90 days", "Revision event", "0.0 (0.0)", "0.0 (0.0)"),
    ("90 days", "Conversion to THA", "0.0 (0.0)", "16.5 (1.6)"),
    ("90 days", "ORIF reoperation", "0.0 (0.0)", "35.0 (3.5)"),
    ("90 days", "Refracture / periprosthetic fracture", "0.6 (0.1)", "2.5 (0.2)"),
    ("sub", "1 Year"),
    ("row", ["Medical complication", "121.0 (12.1)", "355.0 (35.5)"], True),
    ("row", ["Infection-related event", "18.3 (1.8)", "48.9 (4.9)"], True),
    ("row", ["Dislocation event", "30.6 (3.1)", "8.5 (0.8)"], True),
    ("row", ["Revision event", "13.8 (1.4)", "0.1 (0.0)"], True),
    ("row", ["Conversion to THA", "0.0 (0.0)", "26.1 (2.6)"], True),
    ("row", ["ORIF reoperation", "0.0 (0.0)", "107.6 (10.8)"], True),
    ("row", ["Refracture / periprosthetic fracture", "1.9 (0.2)", "7.5 (0.8)"], True),
    ("sub", "2 Years"),
    ("row", ["Medical complication", "121.0 (12.1)", "355.0 (35.5)"], True),
    ("row", ["Infection-related event", "21.1 (2.1)", "52.8 (5.3)"], True),
    ("row", ["Dislocation event", "49.7 (5.0)", "9.2 (0.9)"], True),
    ("row", ["Revision event", "28.0 (2.8)", "0.4 (0.0)"], True),
    ("row", ["Conversion to THA", "0.0 (0.0)", "39.0 (3.9)"], True),
    ("row", ["ORIF reoperation", "0.0 (0.0)", "152.6 (15.3)"], True),
    ("row", ["Refracture / periprosthetic fracture", "3.7 (0.4)", "7.6 (0.8)"], True),
    ("sub", "5 Years"),
    ("row", ["Medical complication", "121.0 (12.1)", "355.0 (35.5)"], True),
    ("row", ["Infection-related event", "30.1 (3.0)", "64.2 (6.4)"], True),
    ("row", ["Dislocation event", "122.9 (12.3)", "12.8 (1.3)"], True),
    ("row", ["Revision event", "71.6 (7.2)", "2.1 (0.2)"], True),
    ("row", ["Conversion to THA", "0.0 (0.0)", "44.6 (4.5)"], True),
    ("row", ["ORIF reoperation", "0.0 (0.0)", "185.7 (18.6)"], True),
    ("row", ["Refracture / periprosthetic fracture", "8.5 (0.8)", "7.8 (0.8)"], True),
    ("sub", "10 Years"),
    ("row", ["Medical complication", "121.0 (12.1)", "355.0 (35.5)"], True),
    ("row", ["Infection-related event", "47.0 (4.7)", "81.6 (8.2)"], True),
    ("row", ["Dislocation event", "285.3 (28.5)", "21.4 (2.1)"], True),
    ("row", ["Revision event", "152.2 (15.2)", "6.0 (0.6)"], True),
    ("row", ["Conversion to THA", "0.0 (0.0)", "53.1 (5.3)"], True),
    ("row", ["ORIF reoperation", "0.0 (0.0)", "235.3 (23.5)"], True),
    ("row", ["Refracture / periprosthetic fracture", "15.2 (1.5)", "8.2 (0.8)"], True),
]
s4_body = [("sub", "90 Days")] + [("row", [r[1], r[2], r[3]], True) for r in s4_data if r[0] == "90 days"]
for item in s4_data:
    if item[0] != "90 days":
        s4_body.append(item)
jbjs_table(
    "Supplemental Table S4.  Cumulative Modeled Events per 1,000 Simulated Patients, All Horizons",
    ["Event", "THA, n (%)", "ORIF, n (%)"],
    s4_body,
    "The values are given as the expected cumulative number of events per 1,000 patients from model "
    "start through the stated horizon, with the percentage of the starting cohort in parentheses. "
    "These are recurrent expected event counts, not the number of unique patients experiencing at "
    "least one event. THA = total hip arthroplasty, ORIF = open reduction and internal fixation.",
    widths=[4.5, 2.5, 2.5], hsize=8.5, bsize=8)

# =========================================================================
# APPENDIX A1 — State list (was a default, unstyled Word table; rebuilt in
# house style, same content, condensed to fit one page)
# =========================================================================
a1_rows = [
    ("WELL_ORIF", "ORIF-native", "Stable native-hip state after fixation"),
    ("DISLOCATION_ORIF", "ORIF-native", "Native-hip instability/dislocation state"),
    ("ORIF_SSI", "ORIF-native", "Post-fixation surgical site infection state"),
    ("REFRACTURE", "ORIF-native", "Refracture after ORIF"),
    ("ORIF_REOP_TUNNEL", "ORIF-native", "Transient perioperative ORIF reoperation state"),
    ("POST_ORIF_REOP", "ORIF-native", "Short post-reoperation dwell state"),
    ("CONVERSION_THA", "Bridge", "Perioperative conversion arthroplasty state after failed ORIF"),
    ("MEDICAL_COMPLICATION", "Shared", "Nonfatal medical complication state"),
    ("WELL_THA", "THA", "Stable arthroplasty state after index acute THA or recovery"),
    ("DISLOCATION", "THA", "THA dislocation state"),
    ("PJI_EARLY", "THA", "Early periprosthetic joint infection"),
    ("PJI_LATE", "THA", "Late periprosthetic joint infection"),
    ("ASEPTIC_LOOSENING", "THA", "Late aseptic loosening state"),
    ("REV_TUNNEL_ASEPTIC", "Revision", "Perioperative aseptic revision state"),
    ("REV_TUNNEL_SEPTIC", "Revision", "Perioperative septic revision state"),
    ("POST_REV_ASEPTIC", "Revision", "Short post-aseptic-revision dwell state"),
    ("POST_REV_SEPTIC", "Revision", "Short post-septic-revision dwell state"),
    ("WELL_REVISED", "Revision", "Longer-term post-revision arthroplasty state"),
    ("DEATH", "Absorbing", "All-cause mortality state"),
]
jbjs_table(
    "Appendix A1.  Model Health States",
    ["State", "Pathway", "Purpose"],
    [("row", list(r), False) for r in a1_rows],
    "Nineteen mutually exclusive health states, grouped by pathway. THA = total hip arthroplasty; "
    "ORIF = open reduction and internal fixation.",
    widths=[2.3, 1.3, 5.4], hsize=9.5, bsize=9.5)

# =========================================================================
# APPENDIX A2 — Allowed transition logic (landscape; the long "allowed next
# states" lists were the source of the bad wrapping in the unstyled version)
# =========================================================================
new_secA2 = doc.add_section(); set_landscape(new_secA2)
a2_rows = [
    ("WELL_ORIF", "WELL_ORIF, DISLOCATION_ORIF, ORIF_SSI, REFRACTURE, ORIF_REOP_TUNNEL, CONVERSION_THA, MEDICAL_COMPLICATION, DEATH", "ORIF-native branch"),
    ("DISLOCATION_ORIF", "WELL_ORIF, DEATH", "Transient native instability state"),
    ("ORIF_SSI", "WELL_ORIF, ORIF_REOP_TUNNEL, DEATH", "Infection can resolve, require reoperation, or die"),
    ("REFRACTURE", "ORIF_REOP_TUNNEL, DEATH", "Refracture routed to reoperation accounting"),
    ("ORIF_REOP_TUNNEL", "POST_ORIF_REOP, DEATH", "30-day perioperative tunnel"),
    ("POST_ORIF_REOP", "WELL_ORIF, DEATH", "Short dwell before return"),
    ("CONVERSION_THA", "WELL_THA, DEATH", "Bridge from failed ORIF to arthroplasty"),
    ("MEDICAL_COMPLICATION", "WELL_ORIF or WELL_THA, DEATH", "Returns to originating well state"),
    ("WELL_THA", "WELL_THA, DISLOCATION, PJI_EARLY, PJI_LATE, ASEPTIC_LOOSENING, MEDICAL_COMPLICATION, DEATH", "Index or recovered arthroplasty state"),
    ("DISLOCATION", "DISLOCATION, REV_TUNNEL_ASEPTIC, DEATH", "Allows recurrence or revision"),
    ("PJI_EARLY", "WELL_THA, REV_TUNNEL_SEPTIC, DEATH", "Early infection can resolve or require revision"),
    ("PJI_LATE", "PJI_LATE, REV_TUNNEL_SEPTIC, DEATH", "Late infection pathway"),
    ("ASEPTIC_LOOSENING", "REV_TUNNEL_ASEPTIC, DEATH", "Revision-driven late failure"),
    ("REV_TUNNEL_ASEPTIC", "POST_REV_ASEPTIC, DEATH", "Aseptic revision perioperative tunnel"),
    ("REV_TUNNEL_SEPTIC", "POST_REV_SEPTIC, DEATH", "Septic revision perioperative tunnel"),
    ("POST_REV_ASEPTIC", "WELL_REVISED, REV_TUNNEL_ASEPTIC, DEATH", "Re-revision allowed"),
    ("POST_REV_SEPTIC", "WELL_REVISED, REV_TUNNEL_SEPTIC, DEATH", "Re-revision allowed"),
    ("WELL_REVISED", "WELL_REVISED, DISLOCATION, PJI_LATE, DEATH", "Retains elevated late revision risks"),
    ("DEATH", "DEATH", "Absorbing state"),
]
jbjs_table(
    "Appendix A2.  Allowed Transition Logic",
    ["From State", "Allowed Next States", "Notes"],
    [("row", list(r), False) for r in a2_rows],
    "Every row includes DEATH as an allowed transition in every cycle via the age- and "
    "condition-specific mortality schedule (Supplemental Table S1), applied in addition to the "
    "state-specific transitions listed here.",
    widths=[2.0, 5.3, 2.1], hsize=9, bsize=9, page_break=False)

doc.save(OUT)
print("Saved all tables to", OUT)
